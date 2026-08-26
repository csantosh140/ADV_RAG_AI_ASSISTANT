"""Tests for Document Parsers (JSON, HTML, Markdown, TXT)."""

import json
import tempfile
from pathlib import Path
import pytest

from ingestion.parsers.json_parser import parse_json
from ingestion.parsers.html_parser import parse_html
from ingestion.parsers.md_parser import parse_markdown
from ingestion.parsers.txt_parser import parse_txt
from ingestion.pipeline import IngestionPipeline


def test_json_parser():
    sample_data = [
        {"title": "LangGraph Overview", "concept": "Cyclic Graph Orchestration", "stars": 5000},
        {"title": "FAISS Overview", "concept": "Dense Vector Search", "stars": 28000}
    ]
    with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False, encoding="utf-8") as f:
        json.dump(sample_data, f)
        temp_path = f.name

    try:
        sections, total = parse_json(temp_path)
        assert total == 2
        assert len(sections) == 2
        assert "LangGraph" in sections[0]["text"]
        assert sections[0]["section_title"] == "LangGraph Overview"
    finally:
        Path(temp_path).unlink(missing_ok=True)


def test_html_parser():
    html_content = """
    <html>
        <head><title>Test Doc</title></head>
        <body>
            <h1>FlashRank Architecture</h1>
            <p>FlashRank is an ultra-fast cross-encoder reranking library based on ONNX.</p>
            <h2>Performance</h2>
            <p>It achieves sub-millisecond reranking without heavy PyTorch overhead.</p>
        </body>
    </html>
    """
    with tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False, encoding="utf-8") as f:
        f.write(html_content)
        temp_path = f.name

    try:
        sections, total = parse_html(temp_path)
        assert total >= 1
        assert any("FlashRank" in s["text"] for s in sections)
    finally:
        Path(temp_path).unlink(missing_ok=True)


def test_csv_parser():
    csv_content = "Name,Role,Experience\nAlice,ML Engineer,5 Years\nBob,Full Stack,3 Years\nCharlie,DevOps,7 Years\n"
    with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False, encoding="utf-8") as f:
        f.write(csv_content)
        temp_path = f.name

    try:
        from ingestion.parsers.csv_parser import parse_csv
        sections, total = parse_csv(temp_path)
        assert total >= 1
        assert "Alice" in sections[0]["text"]
        assert "ML Engineer" in sections[0]["text"]
    finally:
        Path(temp_path).unlink(missing_ok=True)


def test_docx_parser():
    import docx
    from ingestion.parsers.docx_parser import parse_docx

    doc = docx.Document()
    doc.add_heading("Agentic RAG Architecture", level=1)
    doc.add_paragraph("This document tests docx ingestion for vector databases.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Component"
    table.cell(0, 1).text = "Technology"
    table.cell(1, 0).text = "Reranker"
    table.cell(1, 1).text = "FlashRank"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        temp_path = f.name

    try:
        doc.save(temp_path)
        sections, total = parse_docx(temp_path)
        assert total >= 1
        combined = " ".join(s["text"] for s in sections)
        assert "Agentic RAG" in combined
        assert "FlashRank" in combined
    finally:
        Path(temp_path).unlink(missing_ok=True)


def test_url_parser_mock(monkeypatch):
    from ingestion.parsers.url_parser import parse_url

    class MockResponse:
        status_code = 200
        text = "<html><head><title>Mocked Page</title></head><body><h1>Agentic Systems</h1><p>Autonomous loops with verification.</p></body></html>"
        def raise_for_status(self):
            pass

    monkeypatch.setattr("requests.get", lambda url, headers=None, timeout=15: MockResponse())

    sections, title, total = parse_url("https://example.com/test-page")
    assert title == "Mocked Page"
    assert total >= 1
    assert any("Agentic Systems" in s["text"] for s in sections)


def test_ingestion_pipeline_all_formats():
    pipeline = IngestionPipeline()

    # Test JSON processing
    data = {"topic": "Agentic AI", "description": "Self-correcting retrieval loop"}
    with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False, encoding="utf-8") as f:
        json.dump(data, f)
        json_path = f.name

    try:
        doc_meta, chunks = pipeline.process_file(json_path)
        assert doc_meta.file_type == "json"
        assert len(chunks) >= 1
        assert "Agentic AI" in chunks[0].text
    finally:
        Path(json_path).unlink(missing_ok=True)

